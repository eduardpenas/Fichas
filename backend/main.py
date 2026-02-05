import sys
import os
import shutil
import pandas as pd
import zipfile
import io
import subprocess
import tempfile
from fastapi import FastAPI, UploadFile, File, HTTPException, Body
from fastapi.responses import FileResponse, StreamingResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict, Any


sys.path.append(os.path.join(os.path.dirname(__file__), '..', 'src'))

# Ahora sí podemos importar tus scripts mágicos
from procesar_anexo import procesar_anexo
from procesar_cvs import procesar_cvs
from logica_fichas import generar_ficha_2_1, generar_ficha_2_2
from validador import ValidadorFichas, validar_antes_generar

# Modelos Pydantic
class UpdateDataRequest(BaseModel):
    data: List[Dict[str, Any]]
    cliente_nif: str = None
    proyecto_acronimo: str = None

# Inicializamos la APP (El restaurante)
app = FastAPI(title="Generador de Fichas API", version="1.0")

# --- SEGURIDAD (CORS) ---
# Esto permite que el Frontend (que vivirá en el puerto 5173) 
# pueda hablar con el Backend (que vivirá en el puerto 8000).
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], # En producción esto se cambia, pero para desarrollo: ¡barra libre!
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Definimos rutas (Directorios)
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
INPUT_DIR = os.path.join(BASE_DIR, 'inputs')
PROYECTOS_DIR = os.path.join(BASE_DIR, 'proyectos')
os.makedirs(PROYECTOS_DIR, exist_ok=True)

def get_client_dir(client_nif: str):
    """Obtiene la carpeta del cliente, creándola si no existe."""
    client_dir = os.path.join(PROYECTOS_DIR, f"Cliente_{client_nif}")
    os.makedirs(client_dir, exist_ok=True)
    os.makedirs(os.path.join(client_dir, 'data'), exist_ok=True)
    os.makedirs(os.path.join(client_dir, 'history'), exist_ok=True)
    return client_dir

def get_project_dir(client_nif: str, proyecto_acronimo: str):
    """Obtiene la carpeta del proyecto dentro del cliente, creándola si no existe."""
    client_dir = os.path.join(PROYECTOS_DIR, f"Cliente_{client_nif}")
    project_dir = os.path.join(client_dir, proyecto_acronimo)
    os.makedirs(project_dir, exist_ok=True)
    os.makedirs(os.path.join(project_dir, 'data'), exist_ok=True)
    os.makedirs(os.path.join(project_dir, 'history'), exist_ok=True)
    return project_dir

def save_to_history(client_dir: str, data_type: str, data: List[Dict]):
    """Guarda una copia con timestamp en el historial."""
    from datetime import datetime
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    history_file = os.path.join(client_dir, 'history', f"{data_type}_{timestamp}.json")
    df = pd.DataFrame(data)
    df.to_json(history_file, orient='records', force_ascii=False, date_format='iso')

def save_client_name(client_nif: str, nombre: str):
    """Guarda el nombre del cliente en un archivo de configuración."""
    client_dir = get_client_dir(client_nif)
    config_file = os.path.join(client_dir, 'config.json')
    import json
    config = {'nombre': nombre}
    with open(config_file, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=2)
    print(f"💾 SAVE_CLIENT_NAME: NIF={client_nif}, nombre='{nombre}', archivo={config_file}")

def get_client_name(client_nif: str):
    """Obtiene el nombre del cliente desde el archivo de configuración."""
    client_dir = os.path.join(PROYECTOS_DIR, f"Cliente_{client_nif}")
    config_file = os.path.join(client_dir, 'config.json')
    
    print(f"🔍 GET_CLIENT_NAME: NIF={client_nif}")
    print(f"   📁 client_dir={client_dir}")
    print(f"   📄 config_file={config_file}")
    print(f"   ✓ config_file existe={os.path.exists(config_file)}")
    
    # Primero intenta leer el nombre desde config.json
    if os.path.exists(config_file):
        try:
            import json
            with open(config_file, 'r', encoding='utf-8') as f:
                config = json.load(f)
                print(f"   ✓ config.json cargado: {config}")
                if 'nombre' in config and config['nombre']:
                    nombre = config['nombre']
                    print(f"   ✅ Nombre encontrado en config.json: '{nombre}'")
                    return nombre
        except Exception as e:
            print(f"   ❌ Error leyendo config.json: {e}")
    
    # Si no existe nombre en config, intenta obtenerlo del personal.json
    personal_file = os.path.join(client_dir, 'data', 'personal.json')
    print(f"   📄 personal_file={personal_file}, existe={os.path.exists(personal_file)}")
    
    if os.path.exists(personal_file):
        try:
            df = pd.read_json(personal_file)
            print(f"   ✓ personal.json cargado ({len(df)} registros)")
            if not df.empty and 'Nombre' in df.columns:
                nombre = df.iloc[0].get('Nombre', client_nif)
                print(f"   ✅ Nombre encontrado en personal.json: '{nombre}'")
                return nombre
        except Exception as e:
            print(f"   ❌ Error leyendo personal.json: {e}")
    
    # Si no encuentra nada, retorna el NIF
    print(f"   ⚠️ No se encontró nombre, usando NIF: '{client_nif}'")
    return client_nif

# --- ENDPOINTS (LOS PLATOS DE LA CARTA) ---

@app.get("/")
def read_root():
    """Para probar si la API está viva."""
    return {"mensaje": "¡Hola! La API de Fichas está funcionando 🚀"}

@app.get("/clientes")
def list_clients():
    """Lista todos los clientes (carpetas en /proyectos)."""
    try:
        print(f"\n📋 LIST_CLIENTS ENDPOINT")
        
        if not os.path.exists(PROYECTOS_DIR):
            print(f"   ⚠️ PROYECTOS_DIR no existe: {PROYECTOS_DIR}")
            return {"clientes": []}
        
        print(f"   📁 PROYECTOS_DIR: {PROYECTOS_DIR}")
        folders = os.listdir(PROYECTOS_DIR)
        print(f"   📂 Carpetas encontradas: {folders}")
        
        clientes = []
        for folder in folders:
            if folder.startswith("Cliente_"):
                nif = folder.replace("Cliente_", "")
                print(f"   🔍 Procesando cliente: {nif}")
                nombre = get_client_name(nif)
                print(f"      ✅ Nombre obtenido: '{nombre}'")
                
                clientes.append({
                    "nif": nif,
                    "nombre": nombre,
                    "folder": folder
                })
        
        print(f"   ✅ Total de clientes: {len(clientes)}")
        return {"clientes": sorted(clientes, key=lambda x: x['nif'])}
    except Exception as e:
        print(f"   ❌ Error en list_clients: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/clientes")
def create_client(nif: str, nombre: str = None):
    """Crea una nueva carpeta de cliente."""
    try:
        print(f"\n📝 CREATE_CLIENT ENDPOINT")
        print(f"   📌 NIF: {nif}")
        print(f"   📌 Nombre: {nombre or 'No proporcionado'}")
        
        nif = nif.strip().upper()
        
        # Validar que el NIF no esté vacío
        if not nif:
            raise HTTPException(status_code=400, detail="NIF is required")
        
        client_dir = os.path.join(PROYECTOS_DIR, f"Cliente_{nif}")
        
        # Verificar si el cliente ya existe
        if os.path.exists(client_dir):
            print(f"   ⚠️ El cliente ya existe: {client_dir}")
            raise HTTPException(status_code=409, detail=f"Client {nif} already exists")
        
        # Crear la carpeta del cliente
        os.makedirs(client_dir, exist_ok=True)
        print(f"   ✅ Carpeta de cliente creada: {client_dir}")
        
        # Guardar el nombre si se proporciona
        if nombre and nombre.strip():
            try:
                save_client_name(nif, nombre.strip())
                print(f"   ✅ Nombre del cliente guardado: '{nombre.strip()}'")
            except Exception as e:
                print(f"   ⚠️ No se pudo guardar el nombre: {e}")
        
        return {
            "message": f"Client {nif} created successfully",
            "nif": nif,
            "nombre": nombre or nif,
            "folder": f"Cliente_{nif}"
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"   ❌ Error en create_client: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/clientes/{cliente_nif}/proyectos")
def list_proyectos(cliente_nif: str):
    """Lista todos los proyectos de un cliente."""
    try:
        cliente_nif = cliente_nif.strip()
        client_dir = os.path.join(PROYECTOS_DIR, f"Cliente_{cliente_nif}")
        
        if not os.path.exists(client_dir):
            return {"proyectos": []}
        
        proyectos = []
        for folder in os.listdir(client_dir):
            folder_path = os.path.join(client_dir, folder)
            # Excluir carpetas 'data' y 'history'
            if os.path.isdir(folder_path) and folder not in ['data', 'history']:
                proyectos.append({
                    "acronimo": folder,
                    "path": folder_path
                })
        
        return {"proyectos": sorted(proyectos, key=lambda x: x['acronimo'])}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/clientes/{cliente_nif}/nombre")
def set_client_name(cliente_nif: str, nombre: str = None):
    """
    Guarda el nombre de un cliente.
    """
    try:
        cliente_nif = cliente_nif.strip()
        print(f"\n📝 SET_CLIENT_NAME ENDPOINT: NIF={cliente_nif}, nombre_param='{nombre}'")
        
        if not nombre:
            print(f"   ❌ Parámetro nombre vacío")
            raise HTTPException(status_code=400, detail="El nombre no puede estar vacío")
        
        nombre = nombre.strip()
        print(f"   📝 Guardando nombre: '{nombre}'")
        save_client_name(cliente_nif, nombre)
        print(f"   ✅ Nombre guardado exitosamente")
        
        return {
            "status": "success",
            "message": f"Nombre guardado: {nombre}",
            "nif": cliente_nif,
            "nombre": nombre
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"   ❌ Error guardando nombre: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/clientes/{cliente_nif}/proyectos")
def create_proyecto(cliente_nif: str, proyecto_acronimo: str = None):
    """Crea un nuevo proyecto para un cliente."""
    try:
        cliente_nif = cliente_nif.strip()
        
        if not proyecto_acronimo:
            raise HTTPException(status_code=400, detail="proyecto_acronimo es requerido")
        
        proyecto_acronimo = proyecto_acronimo.strip().upper()
        
        print(f"\n✨ CREANDO PROYECTO: {cliente_nif} / {proyecto_acronimo}")
        
        # Crear directorios del proyecto
        project_dir = get_project_dir(cliente_nif, proyecto_acronimo)
        
        print(f"✅ Proyecto creado: {project_dir}")
        
        return {
            "status": "success",
            "message": f"Proyecto {proyecto_acronimo} creado",
            "cliente_nif": cliente_nif,
            "proyecto_acronimo": proyecto_acronimo,
            "path": project_dir
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/clientes/{cliente_nif}")
def delete_client(cliente_nif: str):
    """
    Elimina completamente un cliente y todos sus datos.
    Nota: Esta acción es irreversible.
    """
    try:
        # Limpiar cliente_nif
        cliente_nif = cliente_nif.strip()
        
        client_dir = os.path.join(PROYECTOS_DIR, f"Cliente_{cliente_nif}")
        
        if not os.path.exists(client_dir):
            raise HTTPException(status_code=404, detail=f"Cliente {cliente_nif} no encontrado")
        
        print(f"\n🗑️  ELIMINANDO CLIENTE: {cliente_nif}")
        print(f"📁 Eliminando carpeta: {client_dir}")
        
        # Eliminar la carpeta completa del cliente
        shutil.rmtree(client_dir)
        
        print(f"✅ Cliente {cliente_nif} eliminado correctamente")
        return {
            "status": "success",
            "message": f"Cliente {cliente_nif} y todos sus datos han sido eliminados"
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error eliminando cliente: {e}")
        raise HTTPException(status_code=500, detail=f"Error eliminando cliente: {str(e)}")

@app.post("/upload-anexo")
async def upload_anexo(file: UploadFile = File(...), cliente_nif: str = None, proyecto_acronimo: str = None):
    """
    1. Recibe el archivo Anexo II y opcionalmente cliente_nif + proyecto_acronimo (como parámetros query).
    2. Lo guarda en la carpeta inputs.
    3. Ejecuta tu script 'procesar_anexo.py'.
    4. Si se proporciona cliente_nif y proyecto_acronimo, los JSONs se guardan en la carpeta del proyecto.
    """
    try:
        # Limpiar parámetros
        if cliente_nif:
            cliente_nif = cliente_nif.strip()
        if proyecto_acronimo:
            proyecto_acronimo = proyecto_acronimo.strip().upper()
        
        print(f"\n{'='*60}")
        print(f"📤 UPLOAD-ANEXO INICIADO")
        print(f"{'='*60}")
        print(f"📌 Archivo recibido: {file.filename}")
        print(f"📌 Cliente NIF: {cliente_nif or 'NONE (INPUT_DIR)'}")
        print(f"📌 Proyecto: {proyecto_acronimo or 'NONE'}")
        print(f"📌 Tipo MIME: {file.content_type}")
        print(f"📌 Input DIR: {INPUT_DIR}")
        
        # Guardar el archivo subido
        file_location = os.path.join(INPUT_DIR, "Anexo_Subido.xlsx")
        print(f"💾 Guardando en: {file_location}")
        
        with open(file_location, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        file_size = os.path.getsize(file_location)
        print(f"✅ Archivo guardado. Tamaño: {file_size} bytes")
        
        # Ejecutar tu lógica de extracción PROCESANDO ESPECÍFICAMENTE EL ARCHIVO SUBIDO
        print(f"🍳 Cocinando: Procesando Anexo_Subido.xlsx...")
        procesar_anexo(archivo_especifico="Anexo_Subido.xlsx", cliente_nif=cliente_nif, proyecto_acronimo=proyecto_acronimo)
        print(f"✅ Anexo procesado exitosamente")
        
        # Verificar archivos generados
        if cliente_nif:
            if proyecto_acronimo:
                project_dir = get_project_dir(cliente_nif, proyecto_acronimo)
                output_dir = os.path.join(project_dir, 'data')
                print(f"📁 Archivos guardados en: {output_dir}")
            else:
                client_dir = get_client_dir(cliente_nif)
                output_dir = os.path.join(client_dir, 'data')
                print(f"📁 Archivos guardados en: {output_dir}")
        else:
            output_dir = INPUT_DIR
            print(f"📁 Archivos guardados en: {output_dir}")
        
        output_files = ['Excel_Personal_2.1.json', 'Excel_Colaboraciones_2.2.json', 'Excel_Facturas_2.2.json']
        for out_file in output_files:
            out_path = os.path.join(output_dir, out_file)
            if os.path.exists(out_path):
                file_size = os.path.getsize(out_path)
                print(f"   ✅ {out_file} ({file_size} bytes)")
            else:
                print(f"   ⚠️ {out_file} NO ENCONTRADO")
        
        print(f"{'='*60}\n")
        return {"status": "success", "message": "Anexo procesado y Excels generados"}
    except Exception as e:
        print(f"❌ ERROR EN UPLOAD-ANEXO: {str(e)}")
        import traceback
        print(traceback.format_exc())
        print(f"{'='*60}\n")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload-cvs")
async def upload_cvs(files: List[UploadFile] = File(...), cliente_nif: str = None, proyecto_acronimo: str = None):
    """
    Recibe múltiples PDFs de CVs y los guarda en la carpeta del proyecto.
    - Si se proporciona cliente_nif + proyecto_acronimo: Guarda en Cliente_{nif}/{proyecto}/cvs/
    - Si solo cliente_nif: Guarda en Cliente_{nif}/cvs/ (compatibilidad hacia atrás)
    - Si nada: Guarda en inputs/cvs (compatibilidad hacia atrás)
    """
    # Limpiar parámetros
    if cliente_nif:
        cliente_nif = cliente_nif.strip()
    if proyecto_acronimo:
        proyecto_acronimo = proyecto_acronimo.strip().upper()
    
    print(f"\n{'='*60}")
    print(f"📥 UPLOAD-CVs INICIADO")
    print(f"{'='*60}")
    print(f"📦 Número de archivos recibidos: {len(files)}")
    print(f"📌 Cliente: {cliente_nif or 'NONE'}")
    print(f"📌 Proyecto: {proyecto_acronimo or 'NONE'}")
    
    # Determinar dónde guardar los CVs
    if cliente_nif and proyecto_acronimo:
        # Modo PROYECTO: Guardar en Cliente_{nif}/{proyecto}/cvs/
        project_dir = get_project_dir(cliente_nif, proyecto_acronimo)
        cvs_dir = os.path.join(project_dir, 'cvs')
        print(f"📁 Modo PROYECTO: Guardando en {cvs_dir}")
    elif cliente_nif:
        # Modo CLIENTE: Guardar en Cliente_{nif}/cvs/
        client_dir = get_client_dir(cliente_nif)
        cvs_dir = os.path.join(client_dir, 'cvs')
        print(f"📁 Modo CLIENTE: Guardando en {cvs_dir}")
    else:
        # Modo INPUT_DIR: Guardar en inputs/cvs/
        cvs_dir = os.path.join(INPUT_DIR, 'cvs')
        print(f"📁 Modo INPUT_DIR: Guardando en {cvs_dir}")
    
    # Crear carpeta si no existe
    if not os.path.exists(cvs_dir):
        os.makedirs(cvs_dir)
        print(f"📁 Creada carpeta CVs: {cvs_dir}")
    
    # IMPORTANTE: Limpiar CVs anteriores del proyecto para evitar mezcla
    if cliente_nif and proyecto_acronimo:
        existing_cvs = [f for f in os.listdir(cvs_dir) if f.endswith('.pdf')]
        if existing_cvs:
            print(f"🗑️ Eliminando {len(existing_cvs)} CVs anteriores del proyecto...")
            for old_cv in existing_cvs:
                try:
                    os.remove(os.path.join(cvs_dir, old_cv))
                    print(f"   ❌ Eliminado: {old_cv}")
                except Exception as e:
                    print(f"   ⚠️ Error al eliminar {old_cv}: {e}")
        
    saved_files = []
    for i, file in enumerate(files):
        print(f"   [{i+1}/{len(files)}] Guardando: {file.filename}")
        file_location = os.path.join(cvs_dir, file.filename)
        try:
            with open(file_location, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            file_size = os.path.getsize(file_location)
            print(f"      ✅ Guardado: {file_size} bytes")
            saved_files.append(file.filename)
        except Exception as e:
            print(f"      ❌ Error: {e}")
    
    print(f"✅ UPLOAD-CVs completado: {len(saved_files)} archivos guardados")
    print(f"{'='*60}\n")
    return {"status": "success", "files": saved_files}

@app.post("/process-cvs")
def trigger_process_cvs(cliente_nif: str = None, proyecto_acronimo: str = None):
    """
    Dispara el script de lectura de CVs.
    Si se proporciona cliente_nif y proyecto_acronimo, procesa CVs para ese proyecto específico.
    Si solo cliente_nif, procesa para ese cliente (compatibilidad hacia atrás).
    """
    print(f"\n{'='*60}")
    print(f"🔄 PROCESS-CVs INICIADO")
    print(f"{'='*60}")
    
    # Limpiar parámetros
    if cliente_nif:
        cliente_nif = cliente_nif.strip()
        print(f"📌 Cliente: {cliente_nif}")
    if proyecto_acronimo:
        proyecto_acronimo = proyecto_acronimo.strip().upper()
        print(f"📌 Proyecto: {proyecto_acronimo}")
    else:
        print(f"📌 Modo: Sin cliente (INPUT_DIR)")
    
    try:
        procesar_cvs(cliente_nif=cliente_nif, proyecto_acronimo=proyecto_acronimo)
        print(f"✅ PROCESS-CVs completado exitosamente")
        print(f"{'='*60}\n")
        return {"status": "success", "message": "CVs leídos e integrados en el Excel"}
    except Exception as e:
        print(f"❌ Error en PROCESS-CVs: {e}")
        import traceback
        print(traceback.format_exc())
        print(f"{'='*60}\n")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/personal")
def get_personal_data(cliente_nif: str = None, proyecto_acronimo: str = None):
    """
    Lee los datos de personal. 
    - Si cliente_nif y proyecto_acronimo se proporcionan, los obtiene de la carpeta del proyecto.
    - Si solo cliente_nif se proporciona, los obtiene de la carpeta del cliente.
    - Si el cliente no tiene datos, devuelve lista vacía (cliente nuevo).
    - Si no se proporciona cliente_nif, obtiene del INPUT_DIR (comportamiento heredado para compatibilidad).
    """
    # Limpiar parámetros
    if cliente_nif:
        cliente_nif = cliente_nif.strip()
    if proyecto_acronimo:
        proyecto_acronimo = proyecto_acronimo.strip().upper()
    
    print(f"\n{'='*60}")
    print(f"📥 GET-PERSONAL INICIADO")
    print(f"{'='*60}")
    print(f"📌 cliente_nif (limpio): {cliente_nif}")
    print(f"📌 proyecto_acronimo: {proyecto_acronimo or 'NONE'}")
    
    if cliente_nif:
        # Si se proporciona proyecto_acronimo, buscar en la carpeta del proyecto
        if proyecto_acronimo:
            project_dir = get_project_dir(cliente_nif, proyecto_acronimo)
            json_path = os.path.join(project_dir, 'data', 'Excel_Personal_2.1.json')
            print(f"📁 Buscando datos del proyecto en: {json_path}")
        else:
            client_dir = get_client_dir(cliente_nif)
            json_path = os.path.join(client_dir, 'data', 'Excel_Personal_2.1.json')
            print(f"📁 Buscando datos del cliente en: {json_path}")
        
        # Si el cliente no tiene datos guardados, devolvemos lista vacía (cliente nuevo)
        if not os.path.exists(json_path):
            print(f"⚠️ Cliente nuevo (sin datos guardados). Devolviendo lista vacía")
            print(f"{'='*60}\n")
            return []
        
        print(f"✅ Archivo encontrado")
        df = pd.read_json(json_path)
    else:
        json_path = os.path.join(INPUT_DIR, "Excel_Personal_2.1.json")
        excel_path = os.path.join(INPUT_DIR, "Excel_Personal_2.1.xlsx")
        print(f"📁 Buscando datos en INPUT_DIR")
        print(f"   JSON: {json_path}")
        print(f"   EXCEL: {excel_path}")
        
        if os.path.exists(json_path):
            print(f"✅ Encontrado: {os.path.basename(json_path)}")
            df = pd.read_json(json_path)
        elif os.path.exists(excel_path):
            print(f"✅ Encontrado: {os.path.basename(excel_path)}")
            df = pd.read_excel(excel_path)
        else:
            print(f"❌ NO ENCONTRADO")
            print(f"{'='*60}\n")
            raise HTTPException(status_code=404, detail="No existe Excel ni JSON. Sube el Anexo primero.")
    
    # Rellenamos los NaN (vacíos) con cadenas vacías
    df = df.fillna("")
    
    # Convertimos a lista de diccionarios (JSON)
    datos = df.to_dict(orient="records")
    print(f"📊 Filas devueltas: {len(datos)}")
    if datos:
        print(f"📋 Columnas: {list(datos[0].keys())}")
        print(f"🔍 Primer registro (primeros 3 campos):")
        first_row = datos[0]
        for i, (k, v) in enumerate(list(first_row.items())[:3]):
            print(f"   {k}: {v}")
    print(f"{'='*60}\n")
    return datos

@app.post("/update-personal")
async def update_personal_data(request: UpdateDataRequest):
    """
    Recibe los datos MODIFICADOS desde el Frontend y sobrescribe el archivo.
    - Si cliente_nif + proyecto_acronimo se proporcionan, guarda en Cliente_{nif}/{proyecto}/data/
    - Si solo cliente_nif se proporciona, guarda en Cliente_{nif}/data/ (compatibilidad)
    - Si nada se proporciona, guarda en INPUT_DIR (compatibilidad heredada)
    """
    try:
        # Limpiar parámetros
        if request.cliente_nif:
            request.cliente_nif = request.cliente_nif.strip()
        if request.proyecto_acronimo:
            request.proyecto_acronimo = request.proyecto_acronimo.strip().upper()
        
        df = pd.DataFrame(request.data)
        
        if request.cliente_nif and request.proyecto_acronimo:
            # Modo PROYECTO: Guardar en Cliente_{nif}/{proyecto}/data/
            project_dir = get_project_dir(request.cliente_nif, request.proyecto_acronimo)
            json_path = os.path.join(project_dir, 'data', 'Excel_Personal_2.1.json')
            save_to_history(project_dir, 'personal', request.data)
            print(f"\n{'='*60}")
            print(f"💾 UPDATE-PERSONAL INICIADO (MODO PROYECTO)")
            print(f"{'='*60}")
            print(f"📌 Cliente: {request.cliente_nif}")
            print(f"📌 Proyecto: {request.proyecto_acronimo}")
            print(f"📁 Guardando en: {json_path}")
        elif request.cliente_nif:
            # Modo CLIENTE: Guardar en Cliente_{nif}/data/
            client_dir = get_client_dir(request.cliente_nif)
            json_path = os.path.join(client_dir, 'data', 'Excel_Personal_2.1.json')
            save_to_history(client_dir, 'personal', request.data)
            print(f"\n{'='*60}")
            print(f"💾 UPDATE-PERSONAL INICIADO (MODO CLIENTE)")
            print(f"{'='*60}")
            print(f"📌 Cliente: {request.cliente_nif}")
            print(f"📁 Guardando en: {json_path}")
        else:
            # Comportamiento heredado: guardar en INPUT_DIR
            json_path = os.path.join(INPUT_DIR, "Excel_Personal_2.1.json")
            excel_path = os.path.join(INPUT_DIR, "Excel_Personal_2.1.xlsx")
            print(f"\n{'='*60}")
            print(f"💾 UPDATE-PERSONAL INICIADO (MODO INPUT_DIR)")
            print(f"{'='*60}")
            print(f"📁 Guardando en: {json_path}")
            
            if os.path.exists(json_path):
                formato = "JSON"
            else:
                json_path = excel_path
                formato = "Excel"
        
        df.to_json(json_path, orient='records', force_ascii=False, date_format='iso')
        print(f"✅ Datos guardados correctamente")
        print(f"{'='*60}\n")
        return {"status": "success", "message": f"Datos guardados correctamente"}
    except Exception as e:
        print(f"❌ ERROR: {e}")
        print(f"{'='*60}\n")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/colaboraciones")
def get_colaboraciones_data(cliente_nif: str = None, proyecto_acronimo: str = None):
    """
    Lee los datos de colaboraciones.
    - Si cliente_nif y proyecto_acronimo se proporcionan, los obtiene de la carpeta del proyecto.
    - Si solo cliente_nif se proporciona, los obtiene de la carpeta del cliente.
    - Si el cliente no tiene datos, devuelve lista vacía (cliente nuevo).
    """
    # Limpiar parámetros
    if cliente_nif:
        cliente_nif = cliente_nif.strip()
    if proyecto_acronimo:
        proyecto_acronimo = proyecto_acronimo.strip().upper()
    
    if cliente_nif:
        # Si se proporciona proyecto_acronimo, buscar en la carpeta del proyecto
        if proyecto_acronimo:
            project_dir = get_project_dir(cliente_nif, proyecto_acronimo)
            json_path = os.path.join(project_dir, 'data', 'Excel_Colaboraciones_2.2.json')
        else:
            client_dir = get_client_dir(cliente_nif)
            json_path = os.path.join(client_dir, 'data', 'Excel_Colaboraciones_2.2.json')
        # Si el cliente/proyecto no tiene datos guardados, devolvemos lista vacía
        if not os.path.exists(json_path):
            return []
        df = pd.read_json(json_path)
    else:
        json_path = os.path.join(INPUT_DIR, "Excel_Colaboraciones_2.2.json")
        excel_path = os.path.join(INPUT_DIR, "Excel_Colaboraciones_2.2.xlsx")
        
        if os.path.exists(json_path):
            df = pd.read_json(json_path)
        elif os.path.exists(excel_path):
            df = pd.read_excel(excel_path)
        else:
            raise HTTPException(status_code=404, detail="No existe archivo de Colaboraciones. Sube el Anexo primero.")
    
    df = df.fillna("")
    datos = df.to_dict(orient="records")
    return datos

@app.post("/update-colaboraciones")
async def update_colaboraciones_data(request: UpdateDataRequest):
    """
    Recibe los datos MODIFICADOS de colaboraciones y sobrescribe el archivo.
    - Si cliente_nif + proyecto_acronimo se proporcionan, guarda en Cliente_{nif}/{proyecto}/data/
    - Si solo cliente_nif se proporciona, guarda en Cliente_{nif}/data/ (compatibilidad)
    - Si nada se proporciona, guarda en INPUT_DIR (compatibilidad heredada)
    """
    try:
        # Limpiar parámetros
        if request.cliente_nif:
            request.cliente_nif = request.cliente_nif.strip()
        if request.proyecto_acronimo:
            request.proyecto_acronimo = request.proyecto_acronimo.strip().upper()
        
        df = pd.DataFrame(request.data)
        
        if request.cliente_nif and request.proyecto_acronimo:
            # Modo PROYECTO: Guardar en Cliente_{nif}/{proyecto}/data/
            project_dir = get_project_dir(request.cliente_nif, request.proyecto_acronimo)
            json_path = os.path.join(project_dir, 'data', 'Excel_Colaboraciones_2.2.json')
            save_to_history(project_dir, 'colaboraciones', request.data)
        elif request.cliente_nif:
            # Modo CLIENTE: Guardar en Cliente_{nif}/data/
            client_dir = get_client_dir(request.cliente_nif)
            json_path = os.path.join(client_dir, 'data', 'Excel_Colaboraciones_2.2.json')
            save_to_history(client_dir, 'colaboraciones', request.data)
        else:
            # Comportamiento heredado: guardar en INPUT_DIR
            json_path = os.path.join(INPUT_DIR, "Excel_Colaboraciones_2.2.json")
            excel_path = os.path.join(INPUT_DIR, "Excel_Colaboraciones_2.2.xlsx")
            
            if os.path.exists(json_path):
                formato = "JSON"
            else:
                json_path = excel_path
                formato = "Excel"
        
        df.to_json(json_path, orient='records', force_ascii=False, date_format='iso')
        return {"status": "success", "message": f"Datos guardados correctamente"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/facturas")
def get_facturas_data(cliente_nif: str = None, proyecto_acronimo: str = None):
    """
    Lee los datos de facturas.
    - Si cliente_nif y proyecto_acronimo se proporcionan, los obtiene de la carpeta del proyecto.
    - Si solo cliente_nif se proporciona, los obtiene de la carpeta del cliente.
    - Si el cliente no tiene datos, devuelve lista vacía (cliente nuevo).
    """
    # Limpiar parámetros
    if cliente_nif:
        cliente_nif = cliente_nif.strip()
    if proyecto_acronimo:
        proyecto_acronimo = proyecto_acronimo.strip().upper()
    
    if cliente_nif:
        # Si se proporciona proyecto_acronimo, buscar en la carpeta del proyecto
        if proyecto_acronimo:
            project_dir = get_project_dir(cliente_nif, proyecto_acronimo)
            json_path = os.path.join(project_dir, 'data', 'Excel_Facturas_2.2.json')
        else:
            client_dir = get_client_dir(cliente_nif)
            json_path = os.path.join(client_dir, 'data', 'Excel_Facturas_2.2.json')
        # Si el cliente/proyecto no tiene datos guardados, devolvemos lista vacía
        if not os.path.exists(json_path):
            return []
        df = pd.read_json(json_path)
    else:
        json_path = os.path.join(INPUT_DIR, "Excel_Facturas_2.2.json")
        excel_path = os.path.join(INPUT_DIR, "Excel_Facturas_2.2.xlsx")
        
        if os.path.exists(json_path):
            df = pd.read_json(json_path)
        elif os.path.exists(excel_path):
            df = pd.read_excel(excel_path)
        else:
            raise HTTPException(status_code=404, detail="No existe archivo de Facturas. Sube el Anexo primero.")
    
    df = df.fillna("")
    datos = df.to_dict(orient="records")
    return datos

@app.get("/metadata")
def get_metadata(cliente_nif: str = None):
    """
    Obtiene los metadatos del cliente (entidad solicitante, NIF cliente, año fiscal).
    Estos datos se extraen del Anexo II cuando se procesa.
    """
    import json
    
    if not cliente_nif:
        raise HTTPException(status_code=400, detail="cliente_nif es requerido")
    
    # Limpiar cliente_nif
    cliente_nif = cliente_nif.strip()
    
    client_dir = get_client_dir(cliente_nif)
    metadata_path = os.path.join(client_dir, 'data', 'metadata.json')
    
    print(f"\n📥 GET-METADATA para cliente: {cliente_nif}")
    print(f"📁 Buscando metadata en: {metadata_path}")
    
    if not os.path.exists(metadata_path):
        print(f"⚠️ No existe metadata (cliente nuevo o sin Anexo procesado)")
        return {
            "entidad_solicitante": "",
            "nif_cliente": cliente_nif,
            "anio_fiscal": pd.Timestamp.now().year
        }
    
    try:
        with open(metadata_path, 'r', encoding='utf-8') as f:
            metadata = json.load(f)
        print(f"✅ Metadata cargada correctamente")
        return metadata
    except Exception as e:
        print(f"❌ Error leyendo metadata: {e}")
        raise HTTPException(status_code=500, detail=f"Error leyendo metadata: {str(e)}")

@app.post("/metadata")
def save_metadata(cliente_nif: str = None, entidad_solicitante: str = None, 
                  nif_cliente: str = None, anio_fiscal: int = None):
    """
    Guarda o actualiza los metadatos del cliente.
    """
    import json
    
    if not cliente_nif:
        raise HTTPException(status_code=400, detail="cliente_nif es requerido")
    
    # Limpiar cliente_nif
    cliente_nif = cliente_nif.strip()
    
    print(f"\n📝 SAVE-METADATA para cliente: {cliente_nif}")
    print(f"   📌 Entidad solicitante: {entidad_solicitante}")
    print(f"   📌 NIF cliente: {nif_cliente}")
    print(f"   📌 Año fiscal: {anio_fiscal}")
    
    try:
        client_dir = get_client_dir(cliente_nif)
        metadata_path = os.path.join(client_dir, 'data', 'metadata.json')
        
        # Crear objeto de metadata
        metadata = {
            "entidad_solicitante": entidad_solicitante or "",
            "nif_cliente": nif_cliente or cliente_nif,
            "anio_fiscal": anio_fiscal or pd.Timestamp.now().year
        }
        
        # Guardar
        with open(metadata_path, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)
        
        print(f"✅ Metadata guardada correctamente")
        return {"status": "success", "message": "Metadatos guardados", "metadata": metadata}
    except Exception as e:
        print(f"❌ Error guardando metadata: {e}")
        raise HTTPException(status_code=500, detail=f"Error guardando metadata: {str(e)}")

@app.post("/update-facturas")
async def update_facturas_data(request: UpdateDataRequest):
    """
    Recibe los datos MODIFICADOS de facturas y sobrescribe el archivo.
    - Si cliente_nif + proyecto_acronimo se proporcionan, guarda en Cliente_{nif}/{proyecto}/data/
    - Si solo cliente_nif se proporciona, guarda en Cliente_{nif}/data/ (compatibilidad)
    - Si nada se proporciona, guarda en INPUT_DIR (compatibilidad heredada)
    """
    try:
        # Limpiar parámetros
        if request.cliente_nif:
            request.cliente_nif = request.cliente_nif.strip()
        if request.proyecto_acronimo:
            request.proyecto_acronimo = request.proyecto_acronimo.strip().upper()
        
        df = pd.DataFrame(request.data)
        
        if request.cliente_nif and request.proyecto_acronimo:
            # Modo PROYECTO: Guardar en Cliente_{nif}/{proyecto}/data/
            project_dir = get_project_dir(request.cliente_nif, request.proyecto_acronimo)
            json_path = os.path.join(project_dir, 'data', 'Excel_Facturas_2.2.json')
            save_to_history(project_dir, 'facturas', request.data)
        elif request.cliente_nif:
            # Modo CLIENTE: Guardar en Cliente_{nif}/data/
            client_dir = get_client_dir(request.cliente_nif)
            json_path = os.path.join(client_dir, 'data', 'Excel_Facturas_2.2.json')
            save_to_history(client_dir, 'facturas', request.data)
        else:
            # Comportamiento heredado: guardar en INPUT_DIR
            json_path = os.path.join(INPUT_DIR, "Excel_Facturas_2.2.json")
            excel_path = os.path.join(INPUT_DIR, "Excel_Facturas_2.2.xlsx")
            
            if os.path.exists(json_path):
                formato = "JSON"
            else:
                json_path = excel_path
                formato = "Excel"
        
        df.to_json(json_path, orient='records', force_ascii=False, date_format='iso')
        return {"status": "success", "message": f"Datos guardados correctamente"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/check-available-fichas")
def check_available_fichas(cliente_nif: str = None, proyecto_acronimo: str = None):
    """
    Verifica qué fichas están disponibles para descargar basándose en los JSONs actuales.
    Retorna información sobre qué datos existen sin generar fichas.
    """
    try:
        # Limpiar parámetros
        if cliente_nif:
            cliente_nif = cliente_nif.strip()
        if proyecto_acronimo:
            proyecto_acronimo = proyecto_acronimo.strip().upper()
        
        # Determinar directorio de datos
        if cliente_nif:
            if proyecto_acronimo:
                project_dir = get_project_dir(cliente_nif, proyecto_acronimo)
                data_dir = os.path.join(project_dir, 'data')
            else:
                data_dir = os.path.join(get_client_dir(cliente_nif), 'data')
        else:
            data_dir = INPUT_DIR
        
        # Rutas de JSONs
        json_personal = os.path.join(data_dir, "Excel_Personal_2.1.json")
        json_colaboraciones = os.path.join(data_dir, "Excel_Colaboraciones_2.2.json")
        json_facturas = os.path.join(data_dir, "Excel_Facturas_2.2.json")
        
        # Contar datos en cada JSON
        personal_count = 0
        colaboraciones_count = 0
        facturas_count = 0
        
        if os.path.exists(json_personal):
            try:
                df_personal = pd.read_json(json_personal)
                personal_count = len(df_personal)
            except:
                pass
        
        if os.path.exists(json_colaboraciones):
            try:
                df_colab = pd.read_json(json_colaboraciones)
                colaboraciones_count = len(df_colab)
            except:
                pass
        
        if os.path.exists(json_facturas):
            try:
                df_fact = pd.read_json(json_facturas)
                facturas_count = len(df_fact)
            except:
                pass
        
        # Determinar qué fichas se pueden generar
        puede_generar_2_1 = personal_count > 0
        puede_generar_2_2 = colaboraciones_count > 0 and facturas_count > 0
        
        return {
            "status": "success",
            "puede_generar_2_1": puede_generar_2_1,
            "puede_generar_2_2": puede_generar_2_2,
            "datos": {
                "personal": personal_count,
                "colaboraciones": colaboraciones_count,
                "facturas": facturas_count
            }
        }
    except Exception as e:
        print(f"❌ Error en check_available_fichas: {e}")
        return {
            "status": "error",
            "puede_generar_2_1": False,
            "puede_generar_2_2": False,
            "datos": {
                "personal": 0,
                "colaboraciones": 0,
                "facturas": 0
            }
        }


@app.post("/generate-fichas")
def generate_fichas(cliente_nif: str = None, proyecto_acronimo: str = None, payload: Dict[str, Any] = Body(None)):
    """
    Genera las fichas Word (2.1 y 2.2) usando las plantillas y JSONs.
    - Si se proporciona cliente_nif y proyecto_acronimo, usa los datos del proyecto específico.
    - Si solo cliente_nif, usa los datos del cliente (compatibilidad hacia atrás).
    - Si no se proporciona cliente_nif, usa los datos del INPUT_DIR (compatibilidad hacia atrás).
    Retorna información sobre qué fichas se pueden generar y por qué.
    """
    try:
        # Limpiar parámetros
        if cliente_nif:
            cliente_nif = cliente_nif.strip()
        if proyecto_acronimo:
            proyecto_acronimo = proyecto_acronimo.strip().upper()
        
        # Determinar directorio de datos
        if cliente_nif:
            if proyecto_acronimo:
                print(f"\n📄 GENERATE-FICHAS para cliente: {cliente_nif} / proyecto: {proyecto_acronimo}")
                project_dir = get_project_dir(cliente_nif, proyecto_acronimo)
                data_dir = os.path.join(project_dir, 'data')
            else:
                print(f"\n📄 GENERATE-FICHAS para cliente: {cliente_nif}")
                data_dir = os.path.join(get_client_dir(cliente_nif), 'data')
        else:
            print(f"\n📄 GENERATE-FICHAS usando INPUT_DIR (sin cliente_nif)")
            data_dir = INPUT_DIR
        
        output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'outputs')
        os.makedirs(output_dir, exist_ok=True)
        
        # Rutas de JSONs
        json_personal = os.path.join(data_dir, "Excel_Personal_2.1.json")
        json_colaboraciones = os.path.join(data_dir, "Excel_Colaboraciones_2.2.json")
        json_facturas = os.path.join(data_dir, "Excel_Facturas_2.2.json")
        
        # Plantillas
        plantilla_2_1 = os.path.join(INPUT_DIR, "2.1.docx")
        plantilla_2_2 = os.path.join(INPUT_DIR, "2.2.docx")
        
        # Salidas
        salida_2_1 = os.path.join(output_dir, "Ficha_2_1.docx")
        salida_2_2 = os.path.join(output_dir, "Ficha_2_2.docx")
        
        generadas = []
        errores = []
        avisos = []
        
        # Verificar disponibilidad de datos
        tiene_personal = os.path.exists(json_personal)
        tiene_colaboraciones = os.path.exists(json_colaboraciones)
        tiene_facturas = os.path.exists(json_facturas)
        
        # Verificar si hay datos dentro de los JSONs (no vacíos)
        personal_count = 0
        colaboraciones_count = 0
        facturas_count = 0
        
        if tiene_personal:
            try:
                df_personal = pd.read_json(json_personal)
                personal_count = len(df_personal)
                if personal_count == 0:
                    avisos.append("Ficha 2.1: No hay registros de personal")
            except:
                pass
        
        if tiene_colaboraciones:
            try:
                df_colab = pd.read_json(json_colaboraciones)
                colaboraciones_count = len(df_colab)
            except:
                pass
        
        if tiene_facturas:
            try:
                df_fact = pd.read_json(json_facturas)
                facturas_count = len(df_fact)
            except:
                pass
        
        # Extraer año fiscal del payload
        anio_fiscal = 2024  # por defecto
        if payload and 'anio_fiscal' in payload:
            anio_fiscal = payload.get('anio_fiscal', 2024)
        
        # Generar Ficha 2.1 (solo requiere personal)
        if tiene_personal and personal_count > 0 and os.path.exists(plantilla_2_1):
            try:
                generar_ficha_2_1(json_personal, plantilla_2_1, salida_2_1, anio_fiscal, 'ACR')
                generadas.append("Ficha_2_1.docx")
                print(f"✅ Ficha 2.1 generada ({personal_count} personas)")
            except Exception as e:
                errores.append(f"Error en Ficha 2.1: {str(e)}")
                print(f"❌ Error en Ficha 2.1: {e}")
        elif not tiene_personal or personal_count == 0:
            avisos.append("Ficha 2.1: No hay datos de personal. Cargue un Anexo primero.")
        
        # Generar Ficha 2.2 (requiere colaboraciones y facturas)
        if tiene_colaboraciones and tiene_facturas and os.path.exists(plantilla_2_2):
            try:
                cliente_nombre = None
                cliente_nif_val = None
                if payload:
                    cliente_nombre = payload.get('cliente_nombre') or payload.get('entidad_solicitante')
                    cliente_nif_val = payload.get('cliente_nif') or payload.get('nif_cliente')

                generar_ficha_2_2(json_colaboraciones, json_facturas, plantilla_2_2, salida_2_2, cliente_nombre=cliente_nombre, cliente_nif=cliente_nif_val, anio=anio_fiscal)
                generadas.append("Ficha_2_2.docx")
                print(f"✅ Ficha 2.2 generada ({colaboraciones_count} colaboraciones, {facturas_count} facturas)")
            except Exception as e:
                errores.append(f"Error en Ficha 2.2: {str(e)}")
                print(f"❌ Error en Ficha 2.2: {e}")
        elif not tiene_colaboraciones or not tiene_facturas:
            avisos.append("Ficha 2.2: No hay datos de colaboraciones o facturas.")
        
        if not generadas:
            raise HTTPException(status_code=400, detail=" | ".join(avisos or errores or ["No se puede generar ninguna ficha"]))
        
        return {
            "status": "success",
            "message": f"Fichas generadas: {', '.join(generadas)}",
            "files": generadas,
            "avisos": avisos,
            "puede_generar_2_1": "Ficha_2_1.docx" in generadas,
            "puede_generar_2_2": "Ficha_2_2.docx" in generadas,
            "datos": {
                "personal": personal_count,
                "colaboraciones": colaboraciones_count,
                "facturas": facturas_count
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error en generate_fichas: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/generate-ficha-2-1-only")
def generate_ficha_2_1_only(cliente_nif: str = None, proyecto_acronimo: str = None, payload: Dict[str, Any] = Body(None)):
    """
    Genera solo la Ficha 2.1 (personal).
    Retorna 200 con estatus indicando si se pudo generar o por qué no.
    """
    try:
        # Limpiar parámetros
        if cliente_nif:
            cliente_nif = cliente_nif.strip()
        if proyecto_acronimo:
            proyecto_acronimo = proyecto_acronimo.strip().upper()
        
        # Determinar directorio de datos
        if cliente_nif:
            if proyecto_acronimo:
                project_dir = get_project_dir(cliente_nif, proyecto_acronimo)
                data_dir = os.path.join(project_dir, 'data')
            else:
                data_dir = os.path.join(get_client_dir(cliente_nif), 'data')
        else:
            data_dir = INPUT_DIR
        
        output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'outputs')
        os.makedirs(output_dir, exist_ok=True)
        
        json_personal = os.path.join(data_dir, "Excel_Personal_2.1.json")
        plantilla_2_1 = os.path.join(INPUT_DIR, "2.1.docx")
        salida_2_1 = os.path.join(output_dir, "Ficha_2_1.docx")
        
        # Verificar datos
        if not os.path.exists(json_personal):
            return {
                "success": False,
                "status": "error",
                "message": "❌ No se puede generar Ficha 2.1",
                "aviso": "No hay datos de personal. Cargue un Anexo primero.",
                "file": None
            }
        
        try:
            df_personal = pd.read_json(json_personal)
            if len(df_personal) == 0:
                return {
                    "success": False,
                    "status": "error",
                    "message": "❌ No se puede generar Ficha 2.1",
                    "aviso": "No hay registros de personal.",
                    "file": None
                }
        except Exception as e:
            return {
                "success": False,
                "status": "error",
                "message": "❌ No se puede generar Ficha 2.1",
                "aviso": "Archivo de personal inválido o corrupto.",
                "file": None
            }
        
        # Extraer año fiscal del payload
        anio_fiscal = 2024
        if payload and 'anio_fiscal' in payload:
            anio_fiscal = payload.get('anio_fiscal', 2024)
        
        # Generar Ficha 2.1
        if os.path.exists(plantilla_2_1):
            try:
                generar_ficha_2_1(json_personal, plantilla_2_1, salida_2_1, anio_fiscal, 'ACR')
                print(f"✅ Ficha 2.1 generada ({len(df_personal)} personas)")
                return {
                    "success": True,
                    "status": "success",
                    "message": f"✅ Ficha 2.1 generada ({len(df_personal)} personas)",
                    "aviso": None,
                    "file": "Ficha_2_1.docx"
                }
            except Exception as e:
                print(f"❌ Error al generar Ficha 2.1: {e}")
                return {
                    "success": False,
                    "status": "error",
                    "message": "❌ Error al generar Ficha 2.1",
                    "aviso": f"Error técnico: {str(e)}",
                    "file": None
                }
        else:
            return {
                "success": False,
                "status": "error",
                "message": "❌ No se puede generar Ficha 2.1",
                "aviso": "Plantilla de Ficha 2.1 no encontrada.",
                "file": None
            }
    except Exception as e:
        print(f"❌ Error en generate_ficha_2_1_only: {e}")
        return {
            "success": False,
            "status": "error",
            "message": "❌ Error al generar Ficha 2.1",
            "aviso": f"Error inesperado: {str(e)}",
            "file": None
        }


@app.post("/generate-ficha-2-2-only")
def generate_ficha_2_2_only(cliente_nif: str = None, proyecto_acronimo: str = None, payload: Dict[str, Any] = Body(None)):
    """
    Genera solo la Ficha 2.2 (colaboraciones y facturas).
    Retorna 200 con estatus indicando si se pudo generar o por qué no.
    """
    try:
        # Limpiar parámetros
        if cliente_nif:
            cliente_nif = cliente_nif.strip()
        if proyecto_acronimo:
            proyecto_acronimo = proyecto_acronimo.strip().upper()
        
        # Determinar directorio de datos
        if cliente_nif:
            if proyecto_acronimo:
                project_dir = get_project_dir(cliente_nif, proyecto_acronimo)
                data_dir = os.path.join(project_dir, 'data')
            else:
                data_dir = os.path.join(get_client_dir(cliente_nif), 'data')
        else:
            data_dir = INPUT_DIR
        
        output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'outputs')
        os.makedirs(output_dir, exist_ok=True)
        
        json_colaboraciones = os.path.join(data_dir, "Excel_Colaboraciones_2.2.json")
        json_facturas = os.path.join(data_dir, "Excel_Facturas_2.2.json")
        plantilla_2_2 = os.path.join(INPUT_DIR, "2.2.docx")
        salida_2_2 = os.path.join(output_dir, "Ficha_2_2.docx")
        
        # Verificar datos
        if not os.path.exists(json_colaboraciones) or not os.path.exists(json_facturas):
            return {
                "success": False,
                "status": "error",
                "message": "❌ No se puede generar Ficha 2.2",
                "aviso": "No hay datos de colaboraciones o facturas.",
                "file": None
            }
        
        try:
            df_colab = pd.read_json(json_colaboraciones)
            df_fact = pd.read_json(json_facturas)
            if len(df_colab) == 0 and len(df_fact) == 0:
                return {
                    "success": False,
                    "status": "error",
                    "message": "❌ No se puede generar Ficha 2.2",
                    "aviso": "No hay registros de colaboraciones o facturas.",
                    "file": None
                }
        except Exception as e:
            return {
                "success": False,
                "status": "error",
                "message": "❌ No se puede generar Ficha 2.2",
                "aviso": "Archivos de colaboraciones/facturas inválidos o corruptos.",
                "file": None
            }
        
        # Extraer año fiscal del payload
        anio_fiscal = 2024
        if payload and 'anio_fiscal' in payload:
            anio_fiscal = payload.get('anio_fiscal', 2024)
        
        # Generar Ficha 2.2
        if os.path.exists(plantilla_2_2):
            try:
                cliente_nombre = None
                cliente_nif_val = None
                if payload:
                    cliente_nombre = payload.get('cliente_nombre') or payload.get('entidad_solicitante')
                    cliente_nif_val = payload.get('cliente_nif') or payload.get('nif_cliente')

                generar_ficha_2_2(json_colaboraciones, json_facturas, plantilla_2_2, salida_2_2, cliente_nombre=cliente_nombre, cliente_nif=cliente_nif_val, anio=anio_fiscal)
                print(f"✅ Ficha 2.2 generada ({len(df_colab)} colaboraciones, {len(df_fact)} facturas)")
                return {
                    "success": True,
                    "status": "success",
                    "message": f"✅ Ficha 2.2 generada ({len(df_colab)} colaboraciones, {len(df_fact)} facturas)",
                    "aviso": None,
                    "file": "Ficha_2_2.docx"
                }
            except Exception as e:
                print(f"❌ Error al generar Ficha 2.2: {e}")
                return {
                    "success": False,
                    "status": "error",
                    "message": "❌ Error al generar Ficha 2.2",
                    "aviso": f"Error técnico: {str(e)}",
                    "file": None
                }
        else:
            return {
                "success": False,
                "status": "error",
                "message": "❌ No se puede generar Ficha 2.2",
                "aviso": "Plantilla de Ficha 2.2 no encontrada.",
                "file": None
            }
    except Exception as e:
        print(f"❌ Error en generate_ficha_2_2_only: {e}")
        return {
            "success": False,
            "status": "error",
            "message": "❌ Error al generar Ficha 2.2",
            "aviso": f"Error inesperado: {str(e)}",
            "file": None
        }


@app.get("/download-fichas")
async def download_fichas(cliente_nif: str = None):
    """
    Descarga todas las fichas generadas como un ZIP.
    Si se proporciona cliente_nif, usa los datos del cliente específico.
    """
    try:
        if cliente_nif:
            cliente_nif = cliente_nif.strip()
            print(f"\n⬇️ DOWNLOAD-FICHAS para cliente: {cliente_nif}")
        else:
            print(f"\n⬇️ DOWNLOAD-FICHAS usando OUTPUT_DIR general")
        
        output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'outputs')
        
        if not os.path.exists(output_dir):
            raise HTTPException(status_code=404, detail="No hay fichas generadas. Ejecuta /generate-fichas primero.")
        
        # Crear un ZIP en memoria
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for filename in os.listdir(output_dir):
                file_path = os.path.join(output_dir, filename)
                if os.path.isfile(file_path):
                    # Agregar archivo al ZIP
                    zip_file.write(file_path, arcname=filename)
        
        zip_buffer.seek(0)
        
        return StreamingResponse(
            iter([zip_buffer.getvalue()]),
            media_type="application/zip",
            headers={"Content-Disposition": "attachment; filename=fichas.zip"}
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al descargar fichas: {str(e)}")


@app.get("/download-ficha")
def download_ficha(name: str, cliente_nif: str = None):
    """Descarga un fichero individual desde la carpeta outputs.
    Parámetros: 
      - name: nombre del fichero (por ejemplo: Ficha_2_1.docx)
      - cliente_nif: opcional, para usar datos del cliente específico
    """
    try:
        if cliente_nif:
            cliente_nif = cliente_nif.strip()
            print(f"\n⬇️ DOWNLOAD-FICHA para cliente: {cliente_nif} - archivo: {name}")
        else:
            print(f"\n⬇️ DOWNLOAD-FICHA desde OUTPUT_DIR - archivo: {name}")
        
        output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'outputs')
        file_path = os.path.join(output_dir, os.path.basename(name))
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="Fichero no encontrado")
        return FileResponse(path=file_path, filename=os.path.basename(file_path), media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al descargar fichero: {str(e)}")


@app.get("/preview-ficha")
def preview_ficha(name: str, cliente_nif: str = None):
    """Devuelve una previsualización HTML simple del contenido textual de un .docx.
    No realiza conversiones complejas: extrae párrafos y los devuelve en HTML.
    Parámetros:
      - name: nombre del fichero (por ejemplo: Ficha_2_1.docx)
      - cliente_nif: opcional, para usar datos del cliente específico
    """
    try:
        from docx import Document as DocxDocument

        if cliente_nif:
            cliente_nif = cliente_nif.strip()
            print(f"\n👁️ PREVIEW-FICHA para cliente: {cliente_nif} - archivo: {name}")
        else:
            print(f"\n👁️ PREVIEW-FICHA desde OUTPUT_DIR - archivo: {name}")

        output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'outputs')
        file_path = os.path.join(output_dir, os.path.basename(name))
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="Fichero no encontrado")

        doc = DocxDocument(file_path)
        html_parts = ["<div style='font-family:Arial,Helvetica,sans-serif;padding:16px'>"]
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Escape minimal HTML
                safe = (text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))
                html_parts.append(f"<p style='margin:6px 0'>{safe}</p>")
        html_parts.append("</div>")
        html = "\n".join(html_parts)
        return HTMLResponse(content=html, status_code=200)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al generar previsualización: {str(e)}")


@app.get("/preview-ficha-pdf")
def preview_ficha_pdf(name: str):
    """Devuelve el .docx convertido a PDF usando LibreOffice soffice si está disponible.
    Si no, retorna el HTML de fallback.
    """
    try:
        output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'outputs')
        file_path = os.path.join(output_dir, os.path.basename(name))
        
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="Fichero no encontrado")

        # Crear temporal para PDF
        with tempfile.TemporaryDirectory() as tmpdir:
            # Intentar conversión con soffice (LibreOffice)
            try:
                cmd = [
                    "soffice",
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", tmpdir,
                    file_path
                ]
                result = subprocess.run(cmd, capture_output=True, timeout=30, text=True)
                
                if result.returncode != 0:
                    raise Exception(f"soffice error: {result.stderr}")
                
                # El PDF tendrá el nombre del archivo original pero con extensión .pdf
                base_name = os.path.splitext(os.path.basename(file_path))[0]
                pdf_path = os.path.join(tmpdir, f"{base_name}.pdf")
                
                if not os.path.exists(pdf_path):
                    raise Exception(f"PDF no se generó. Buscaba: {pdf_path}")

                # Leer y devolver PDF
                with open(pdf_path, 'rb') as f:
                    pdf_data = f.read()
                
                return FileResponse(
                    io.BytesIO(pdf_data),
                    media_type="application/pdf",
                    filename=f"{os.path.splitext(os.path.basename(name))[0]}.pdf"
                )
            except Exception as e:
                # Si falla, caer a HTML fallback
                print(f"⚠️ Conversión PDF falló: {str(e)}. Usando fallback HTML.")
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                html_parts = ["<div style='font-family:Arial,Helvetica,sans-serif;padding:16px;background:#f9f9f9;border:1px solid #ddd'>"]
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        safe = (text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))
                        html_parts.append(f"<p style='margin:8px 0;line-height:1.6'>{safe}</p>")
                html_parts.append("<p style='margin-top:20px;color:#666;font-size:0.9em;border-top:1px solid #ddd;padding-top:10px'>💡 Previsualización en HTML (LibreOffice no disponible)</p></div>")
                html = "\n".join(html_parts)
                return HTMLResponse(content=html, status_code=200)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al generar PDF: {str(e)}")


@app.post("/validate")
def validate_data(cliente_nif: str = None, proyecto_acronimo: str = None):
    """
    Valida todos los archivos JSON generados.
    - Si cliente_nif y proyecto_acronimo se proporcionan, valida los datos del proyecto específico.
    - Si solo cliente_nif, valida los datos del cliente (compatibilidad hacia atrás).
    - Si no se proporciona cliente_nif, valida los datos del INPUT_DIR (compatibilidad hacia atrás).
    Retorna reporte de validación con errores, advertencias y estado.
    """
    try:
        # Limpiar parámetros
        if cliente_nif:
            cliente_nif = cliente_nif.strip()
        if proyecto_acronimo:
            proyecto_acronimo = proyecto_acronimo.strip().upper()
        
        # Determinar directorio de datos
        if cliente_nif:
            if proyecto_acronimo:
                print(f"\n✅ VALIDATE para cliente: {cliente_nif} / proyecto: {proyecto_acronimo}")
                project_dir = get_project_dir(cliente_nif, proyecto_acronimo)
                data_dir = os.path.join(project_dir, 'data')
            else:
                print(f"\n✅ VALIDATE para cliente: {cliente_nif}")
                data_dir = os.path.join(get_client_dir(cliente_nif), 'data')
        else:
            print(f"\n✅ VALIDATE usando INPUT_DIR (sin cliente_nif)")
            data_dir = INPUT_DIR
        
        archivo_personal = os.path.join(data_dir, "Excel_Personal_2.1.json")
        archivo_colaboraciones = os.path.join(data_dir, "Excel_Colaboraciones_2.2.json")
        archivo_facturas = os.path.join(data_dir, "Excel_Facturas_2.2.json")
        
        print(f"   📁 Leyendo de: {data_dir}")
        
        # Verificar que existan los archivos
        if not os.path.exists(archivo_personal):
            raise HTTPException(status_code=400, detail="Archivo Personal no existe. Ejecute /upload-anexo primero.")
        if not os.path.exists(archivo_colaboraciones):
            raise HTTPException(status_code=400, detail="Archivo Colaboraciones no existe. Ejecute /upload-anexo primero.")
        if not os.path.exists(archivo_facturas):
            raise HTTPException(status_code=400, detail="Archivo Facturas no existe. Ejecute /upload-anexo primero.")
        
        # Ejecutar validación
        es_valido, resumen = validar_antes_generar(
            archivo_personal,
            archivo_colaboraciones,
            archivo_facturas
        )
        
        return {
            "status": "valid" if es_valido else "invalid",
            "exitosa": es_valido,
            "resumen": resumen,
            "puede_generar": es_valido
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error en validación: {str(e)}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)