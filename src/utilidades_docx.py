import pandas as pd
from docx.shared import Pt, Cm, RGBColor
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# --- Funciones Auxiliares ---

def cm_to_pt(cm):
    """Convierte centímetros a puntos."""
    return cm * 28.3465

def formatea_euro(valor, unidad="€"):
    """Formatea un número a moneda europea (1.234,56 €)."""
    try:
        val = float(valor)
        return "{:,.2f}".format(val).replace(",", "X").replace(".", ",").replace("X", ".") + f" {unidad}"
    except:
        return f"0,00 {unidad}"

def get_value_or_default(df, index, column_name, default=""):
    """Obtiene un valor de un DataFrame de forma segura."""
    try:
        val = df.iloc[index][column_name]
        return val if pd.notna(val) else default
    except:
        return default

# --- Funciones de Formato DOCX ---

def set_cell_color(cell, color_hex):
    """Pone color de fondo a una celda."""
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), color_hex)
    cell_properties.append(shading)

# Alias para compatibilidad
set_cell_background = set_cell_color 

def set_text_format(cell, bold=False):
    """Aplica formato estándar (Arial 10) a una celda existente."""
    if not cell.paragraphs:
        cell.add_paragraph()
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.bold = bold
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def add_text_to_cell(cell, text, bold=False):
    """Limpia una celda y añade texto nuevo con formato."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.bold = bold
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def set_cell_format(cell, text, font_name="Arial", font_size=10, bold=True, bg_color="F2F2F2"):
    """Formatea una celda completa: texto, fuente, tamaño y color de fondo."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    set_cell_color(cell, bg_color)

def configurar_tabla_estandar(table):
    """Aplica formato Grid, alineación izquierda y bordes XML a una tabla."""
    table.style = 'Table Grid'
    tbl = table._element
    tblPr = tbl.find(ns.qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    jc = OxmlElement('w:jc')
    jc.set(ns.qn('w:val'), 'left')
    tblPr.append(jc)
    
    # Sangría de -0.83 cm convertida a twips
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(ns.qn('w:w'), str(int(-0.83 * 567))) 
    tblInd.set(ns.qn('w:type'), 'dxa')
    tblPr.append(tblInd)

# --- Funciones Complejas ---

# Ubicación: src/utilidades_docx.py

import pandas as pd
from docx.shared import Pt, Cm, RGBColor
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# ... (Mantén las funciones auxiliares cm_to_pt, formatea_euro, etc. igual que antes) ...
# Asegúrate de que cm_to_pt existe:
def cm_to_pt(cm):
    """Convierte centímetros a puntos."""
    return cm * 28.3465

# ... (Resto de funciones auxiliares) ...

# --- FUNCIÓN MODIFICADA ---

def create_titled_box(doc, title_text, body_text, height=None, width=None):
    """
    Crea una caja con título y cuerpo de texto.
    Permite definir altura y anchura (en Puntos).
    Si no se definen, usa valores por defecto (5.3cm alto, 18.33cm ancho).
    """
    
    # 1. Título
    title = doc.add_paragraph()
    run = title.add_run(title_text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # Sangrías del título
    title.paragraph_format.left_indent = Cm(-1)
    title.paragraph_format.right_indent = Cm(-0.75)
    title.paragraph_format.line_spacing = 1.0
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(0)

    # 2. Tabla contenedora
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    
    # Ajustes XML
    tbl = table._element
    tblPr = tbl.find(ns.qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    jc = OxmlElement('w:jc')
    jc.set(ns.qn('w:val'), 'left')
    tblPr.append(jc)

    # Sangría de tabla (-0.83 cm)
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(ns.qn('w:w'), str(int(-0.83 * 567))) 
    tblInd.set(ns.qn('w:type'), 'dxa')
    tblPr.append(tblInd)

    # --- LÓGICA DE ANCHURA (WIDTH) ---
    tblW = OxmlElement('w:tblW')
    if width:
        # Si viene un width (en Puntos), XML necesita Twips (1 Pt = 20 Twips)
        width_twips = int(width * 20)
        tblW.set(ns.qn('w:w'), str(width_twips))
    else:
        # Valor por defecto (18.33 cm convertido a twips)
        tblW.set(ns.qn('w:w'), str(int(18.33 * 567))) 
        
    tblW.set(ns.qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # --- LÓGICA DE ALTURA (HEIGHT) ---
    if height:
        # Si viene un height (en Puntos), usamos Pt()
        table.rows[0].height = Pt(height)
    else:
        # Valor por defecto (5.3 cm)
        table.rows[0].height = Cm(5.3)

    # 4. Insertar texto dentro
    cell = table.cell(0, 0)
    cell.text = "" 
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(body_text) if body_text else " ")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 0, 0)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    return table

def crear_tabla_coste_colaboracion(doc, facturas_entidad):
    """Genera la tabla de costes sumando facturas (Ficha 2.2)."""
    
    # Función interna formateo (reutilizamos la global para evitar errores)
    def formatea_numero_local(valor):
        try:
            num = float(valor)
            return f"{num:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        except (ValueError, TypeError):
            return str(valor)

    total_importe = 0
    num_facturas = len(facturas_entidad)
    filas_extra = max(num_facturas - 6, 0)
    filas_total = 34 + filas_extra

    table = doc.add_table(rows=filas_total, cols=3)
    table.style = "Table Grid"

    column_widths = [Cm(4.27), Cm(11.24), Cm(2.76)]
    for row in table.rows:
        for j in range(3):
            row.cells[j].width = column_widths[j]
        row.height = Cm(0.3)

    # --- CORRECCIÓN DEL WARNING AQUÍ ---
    tbl = table._element
    tblPr = tbl.find(ns.qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # -----------------------------------

    jc = OxmlElement('w:jc')
    jc.set(ns.qn('w:val'), 'left')
    tblPr.append(jc)

    tblInd = OxmlElement('w:tblInd')
    tblInd.set(ns.qn('w:w'), str(int(-0.83 * 567)))
    tblInd.set(ns.qn('w:type'), 'dxa')
    tblPr.append(tblInd)

    tblW = OxmlElement('w:tblW')
    tblW.set(ns.qn('w:w'), str(int(18.33 * 567)))
    tblW.set(ns.qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # Texto de cabeceras y partidas
    text_cells = {
        (0, 0): "PARTIDA", (0, 1): "CONCEPTO", (0, 2): "IMPORTE (€)",
        (2, 0): "PERSONAL", (8 + filas_extra, 0): "AMORTIZACIÓN DE ACTIVOS MATERIALES E INMATERIALES",
        (14 + filas_extra, 0): "MATERIAL FUNGIBLE", (20 + filas_extra, 0): "COLABORACIONES EXTERNAS",
        (26 + filas_extra, 0): "OTROS GASTOS", (7 + filas_extra, 1): "TOTAL PERSONAL",
        (13 + filas_extra, 1): "TOTAL AMORTIZACIONES", (19 + filas_extra, 1): "TOTAL MATERIAL",
        (25 + filas_extra, 1): "TOTAL COLABORACIONES", (31 + filas_extra, 1): "TOTAL OTROS GASTOS",
        (33 + filas_extra, 0): "TOTAL IMPORTE SUBCONTRATACIÓN DE LA ENTIDAD COLABORADORA"
    }

    # Aplicar formato
    for (row, col), text in text_cells.items():
        cell = table.cell(row, col)
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.bold = True

    # Colorear celdas
    cells_to_color = [(0, 0), (0, 1), (0, 2), (2, 0),
                      (8 + filas_extra, 0), (14 + filas_extra, 0),
                      (20 + filas_extra, 0), (26 + filas_extra, 0),
                      (7 + filas_extra, 1), (13 + filas_extra, 1),
                      (19 + filas_extra, 1), (25 + filas_extra, 1),
                      (31 + filas_extra, 1), (33 + filas_extra, 0)]
    for (row, col) in cells_to_color:
        set_cell_color(table.cell(row, col), "E2EFD9")

    # Fusionar columnas verticales para partidas
    partidas = [(2, 7 + filas_extra), (8 + filas_extra, 13 + filas_extra),
                (14 + filas_extra, 19 + filas_extra), (20 + filas_extra, 25 + filas_extra),
                (26 + filas_extra, 31 + filas_extra)]
    for start, end in partidas:
        table.cell(start, 0).merge(table.cell(end, 0))

    table.cell(1, 0).merge(table.cell(1, 2))
    table.cell(32 + filas_extra, 0).merge(table.cell(32 + filas_extra, 2))
    table.cell(33 + filas_extra, 0).merge(table.cell(33 + filas_extra, 1))

    # Insertar facturas en la sección de PERSONAL (desde fila 2)
    for i, (_, factura) in enumerate(facturas_entidad.iterrows()):
        fila_actual = 2 + i
        table.cell(fila_actual, 1).text = str(factura.get("Nombre factura", ""))
        if table.cell(fila_actual, 1).paragraphs:
            run = table.cell(fila_actual, 1).paragraphs[0].runs[0]
            run.font.name = "Arial"
            run.font.size = Pt(10)

        try:
            val_imp = factura.get("Importe (€)", 0)
            importe_num = float(val_imp)
        except ValueError:
            importe_num = float(str(val_imp).replace(".", "").replace(",", "."))

        total_importe += importe_num

        importe_str = f"{formatea_numero_local(importe_num)} €"
        table.cell(fila_actual, 2).text = importe_str
        if table.cell(fila_actual, 2).paragraphs:
            run = table.cell(fila_actual, 2).paragraphs[0].runs[0]
            run.font.name = "Arial"
            run.font.size = Pt(10)

    # Insertar total en fila TOTAL PERSONAL
    fila_total_personal = 7 + filas_extra
    table.cell(fila_total_personal, 2).text = f"{formatea_numero_local(total_importe)} €"
    if table.cell(fila_total_personal, 2).paragraphs:
        run = table.cell(fila_total_personal, 2).paragraphs[0].runs[0]
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.bold = True

    # Insertar total general
    fila_total_general = 33 + filas_extra
    table.cell(fila_total_general, 2).text = f"{formatea_numero_local(total_importe)} €"
    if table.cell(fila_total_general, 2).paragraphs:
        run = table.cell(fila_total_general, 2).paragraphs[0].runs[0]
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.bold = True

    return table