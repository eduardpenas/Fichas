import pandas as pd
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
import os

# Importamos las utilidades
from utilidades_docx import (
    formatea_euro, get_value_or_default, add_text_to_cell, 
    set_cell_color, set_text_format, create_titled_box, cm_to_pt,
    set_cell_format, crear_tabla_coste_colaboracion
)

# ==========================================
# FICHA 2.1 (Personal)
# ==========================================
def generar_ficha_2_1(ruta_excel, ruta_plantilla_base, ruta_salida_final, anio, acronimus):
    """Genera la Ficha 2.1 replicando exactamente la lógica del notebook."""
    print(f"Leyendo Excel: {ruta_excel}")
    df_ficha = pd.read_excel(ruta_excel)

    # Ordenar por nombre si es necesario
    if not df_ficha['Nombre'].is_monotonic_increasing:
        df_ficha = df_ficha.sort_values(by='Nombre').reset_index(drop=True)

    doc_master = Document()

    for index in range(len(df_ficha)):
        # --- TABLA 1: PERSONAL PARTICIPANTE ---
        title1 = doc_master.add_paragraph()
        run = title1.add_run("1. IDENTIFICACIÓN DE PERSONAL PARTICIPANTE DE ENTIDAD SOLICITANTE:")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        title1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        title1.paragraph_format.left_indent = Cm(-1)
        title1.paragraph_format.right_indent = Cm(-0.75)
        title1.paragraph_format.line_spacing = 1.0

        table1 = doc_master.add_table(rows=8, cols=4)
        table1.style = 'Table Grid'
        
        # XML Tabla 1
        tbl = table1._element
        tblPr = tbl.find(ns.qn('w:tblPr'))
        if tblPr is None: tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
        jc = OxmlElement('w:jc'); jc.set(ns.qn('w:val'), 'left'); tblPr.append(jc)
        tblInd = OxmlElement('w:tblInd'); tblInd.set(ns.qn('w:w'), str(int(-0.83 * 567))); tblInd.set(ns.qn('w:type'), 'dxa'); tblPr.append(tblInd)
        tblW = OxmlElement('w:tblW'); tblW.set(ns.qn('w:w'), str(int(18.33 * 567))); tblW.set(ns.qn('w:type'), 'dxa'); tblPr.append(tblW)

        for row in table1.rows: row.height = Cm(0.5)
        
        column_widths = [Cm(4.32), Cm(6.5), Cm(3.25), Cm(4.5)]
        for row in table1.rows:
            for j, width in enumerate(column_widths):
                row.cells[j].width = width

        # Fusiones y Formato Tabla 1
        table1.rows[0].cells[1].merge(table1.rows[0].cells[3])
        table1.rows[1].cells[1].merge(table1.rows[1].cells[3])
        
        for row in table1.rows:
            set_cell_color(row.cells[0], "F2F2F2")
            set_text_format(row.cells[0], bold=True)
        for fila in [2, 3, 4, 5, 6, 7]:
            set_cell_color(table1.rows[fila].cells[2], "F2F2F2")
            set_text_format(table1.rows[fila].cells[2], bold=True)

        # Rellenar Datos Tabla 1
        fields1 = [
            ("Nombre", get_value_or_default(df_ficha, index, "Nombre"), "", ""),
            ("Apellidos", get_value_or_default(df_ficha, index, "Apellidos"), "", ""),
            ("Coste horario (€/hora)", formatea_euro(get_value_or_default(df_ficha, index, 'Coste horario (€/hora)', 0), "€/h"), "", ""),
            ("Coste total (€)", formatea_euro(get_value_or_default(df_ficha, index, 'Coste total (€)', 0)), "Horas totales", int(get_value_or_default(df_ficha, index, "Horas totales", 0))),
            ("Coste I+D (€)", formatea_euro(get_value_or_default(df_ficha, index, 'Coste I+D (€)', 0)), "Horas I+D", get_value_or_default(df_ficha, index, "Horas I+D")),
            ("Coste IT (€)", formatea_euro(get_value_or_default(df_ficha, index, 'Coste total (€)', 0)), "Horas IT", int(get_value_or_default(df_ficha, index, "Horas totales", 0))),
            ("Departamento", get_value_or_default(df_ficha, index, "Departamento"), "Puesto actual", get_value_or_default(df_ficha, index, "Puesto actual")),
            ("Titulación 1", get_value_or_default(df_ficha, index, "Titulación 1"), "Titulación 2", get_value_or_default(df_ficha, index, "Titulación 2"))
        ]

        for i, row_data in enumerate(fields1):
            row_cells = table1.rows[i].cells
            add_text_to_cell(row_cells[0], row_data[0], bold=True)
            if i < 2:
                add_text_to_cell(row_cells[1], row_data[1])
            else:
                add_text_to_cell(row_cells[1], str(row_data[1]))
                add_text_to_cell(row_cells[2], row_data[2], bold=True)
                add_text_to_cell(row_cells[3], str(row_data[3]))

        doc_master.add_paragraph()

        # =========================================================================
        # TABLA 2: COLABORACIÓN EXTERNA
        # =========================================================================
        title2 = doc_master.add_paragraph()
        run = title2.add_run("2. IDENTIFICACIÓN DE PERSONAL PARTICIPANTE DE COLABORACIÓN EXTERNA:")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        title2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        title2.paragraph_format.left_indent = Cm(-1)
        title2.paragraph_format.right_indent = Cm(-0.75)
        title2.paragraph_format.line_spacing = 1.0

        table2 = doc_master.add_table(rows=8, cols=5)
        table2.style = 'Table Grid'
        
        # XML Tabla 2
        tbl = table2._element
        tblPr = tbl.find(ns.qn('w:tblPr'))
        if tblPr is None: tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
        jc = OxmlElement('w:jc'); jc.set(ns.qn('w:val'), 'left'); tblPr.append(jc)
        tblInd = OxmlElement('w:tblInd'); tblInd.set(ns.qn('w:w'), str(int(-0.83 * 567))); tblInd.set(ns.qn('w:type'), 'dxa'); tblPr.append(tblInd)
        tblW = OxmlElement('w:tblW'); tblW.set(ns.qn('w:w'), str(int(18.33 * 567))); tblW.set(ns.qn('w:type'), 'dxa'); tblPr.append(tblW)

        column_widths_2 = [Cm(4.32), Cm(4.75), Cm(4), Cm(1.76), Cm(2.23)]
        for row in table2.rows:
            for j, width in enumerate(column_widths_2):
                row.cells[j].width = width

        # Fusiones y Estilo Tabla 2
        table2.cell(0, 1).merge(table2.cell(0, 2))
        table2.cell(1, 1).merge(table2.cell(1, 4))
        table2.cell(2, 1).merge(table2.cell(2, 4))
        table2.cell(3, 1).merge(table2.cell(3, 4))
        for r in range(4, 8): table2.cell(r, 3).merge(table2.cell(r, 4))

        for row in table2.rows:
            set_cell_color(row.cells[0], "F2F2F2")
            set_text_format(row.cells[0], bold=True)
        for fila in [4, 5, 6, 7]:
            set_cell_color(table2.rows[fila].cells[2], "F2F2F2")
            set_text_format(table2.rows[fila].cells[2], bold=True)
        set_cell_color(table2.rows[0].cells[-2], "F2F2F2")

        headers_2 = [
            ("Entidad Colaboradora", "", "", "NIF", ""),
            ("Nombre", "", "", "", ""),
            ("Apellidos", "", "", "", ""),
            ("Perfil profesional", "", "", "", ""),
            ("Número de personas", "", "Coste horario (€/hora)", "", ""),
            ("Coste total (€)", "", "Horas totales", "", ""),
            ("Titulación 1", "", "Titulación 2", "", ""),
            ("Titulación 3", "", "Titulación 4", "", "")
        ]

        for row_idx, header in enumerate(headers_2):
            for col_idx, text in enumerate(header):
                cell = table2.cell(row_idx, col_idx)
                if text:
                    cell.text = "" 
                    p = cell.paragraphs[0]
                    run = p.add_run(text)
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                    run.bold = True
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        doc_master.add_paragraph()

        # ==========================================
        # SECCIÓN 3 y 4 (TEXTOS DINÁMICOS COMPLETOS)
        # ==========================================
        
        # 1. Extracción de variables desde el Excel
        nombre = str(get_value_or_default(df_ficha, index, "Nombre", "")).strip()
        apellidos = str(get_value_or_default(df_ficha, index, "Apellidos", "")).strip()
        titulacion = str(get_value_or_default(df_ficha, index, "Titulación 1", "")).strip()
        puesto = str(get_value_or_default(df_ficha, index, "Puesto actual", "")).strip()
        departamento = str(get_value_or_default(df_ficha, index, "Departamento", "")).strip()
        horas = int(get_value_or_default(df_ficha, index, "Horas totales", 0))
        coste_total_val = get_value_or_default(df_ficha, index, "Coste total (€)", 0)
        
        # Variables de Empresas Anteriores (Historial)
        empresa1 = str(get_value_or_default(df_ficha, index, "EMPRESA 1", "")).strip()
        periodo1 = str(get_value_or_default(df_ficha, index, "PERIODO 1", "")).strip()
        cargo1 = str(get_value_or_default(df_ficha, index, "PUESTO 1", "")).strip()
        
        empresa2 = str(get_value_or_default(df_ficha, index, "EMPRESA 2", "")).strip()
        periodo2 = str(get_value_or_default(df_ficha, index, "PERIODO 2", "")).strip()
        cargo2 = str(get_value_or_default(df_ficha, index, "PUESTO 2", "")).strip()
        
        empresa3 = str(get_value_or_default(df_ficha, index, "EMPRESA 3", "")).strip()
        periodo3 = str(get_value_or_default(df_ficha, index, "PERIODO 3", "")).strip()
        cargo3 = str(get_value_or_default(df_ficha, index, "PUESTO 3", "")).strip()

        # 2. Construcción de la frase (Sección 3)
        frase = f"{nombre} {apellidos} ha trabajado en {empresa1} durante el periodo de {periodo1} ocupando el cargo de {cargo1}."
        
        if empresa2:
            frase += f" También trabajó en {empresa2} durante el periodo de {periodo2} ocupando el cargo de {cargo2}."
        if empresa3:
            frase += f" Por último, trabajó en {empresa3} durante el periodo de {periodo3} ocupando el cargo de {cargo3}."

        # 3. Construcción de actividades y costes (Sección 4)
        texto_actividades = ""
        for n in range(1, 5):
            act = str(get_value_or_default(df_ficha, index, f"Actividad {n}", "")).strip()
            if act: 
                texto_actividades += f"{act}\n"

        # Formateo manual del coste para texto (estilo europeo 1.234,56)
        try: 
            coste_total = f"{float(coste_total_val):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        except: 
            coste_total = "0,00"
            
        # Extraer año de inicio del periodo 1 (ej: "2010" de "2010-2015")
        anio_inicio = periodo1[0:4] if len(periodo1) >= 4 else periodo1

        actividad_4_1 = (
            f"{nombre} {apellidos}, con titulación en {titulacion} ocupa el puesto de {puesto} "
            f"dentro del Departamento de {departamento}, empresa de la que forma parte desde {anio_inicio} "
            f"y participa de manera activa durante la ejecución del proyecto. Concretamente participa durante "
            f"{horas} horas en {anio}, lo que supone un gasto de {coste_total} €.\n\n"
            f"Su participación se considera esencial para la correcta ejecución del presente proyecto llevado "
            f"a cabo durante la anualidad {anio}, participando concretamente en las siguientes fases y tareas del mismo:\n\n"
            f"{texto_actividades.strip()}"
        )

        # 4. Inserción en el documento
        doc_master.add_paragraph()
        create_titled_box(
            doc_master, 
            "3. EXPERIENCIA PROFESIONAL RELACIONADA CON LA ACTIVIDAD DESARROLLADA EN EL PROYECTO:", 
            frase,
            height=cm_to_pt(2.28), 
            width=cm_to_pt(18.24)
        )
        
        doc_master.add_paragraph()
        create_titled_box(
            doc_master, 
            "4. FUNCIONES ASIGNADAS/DESARROLLADAS EN EL PROYECTO.", 
            actividad_4_1,
            height=cm_to_pt(2.28), 
            width=cm_to_pt(18.24)
        )

        if index < len(df_ficha) - 1:
            doc_master.add_page_break()

    # Fusión Ficha 2.1
    fusionar_y_guardar(doc_master, ruta_plantilla_base, ruta_salida_final)


# ==========================================
# FICHA 2.2 (Colaboraciones)
# ==========================================
def generar_ficha_2_2(ruta_colaboraciones, ruta_facturas, ruta_plantilla_base, ruta_salida_final):
    """Genera la Ficha 2.2 replicando exactamente la lógica del notebook."""
    print(f"Leyendo Excel Colaboraciones: {ruta_colaboraciones}")
    df_colab = pd.read_excel(ruta_colaboraciones)
    
    print(f"Leyendo Excel Facturas: {ruta_facturas}")
    df_facturas = pd.read_excel(ruta_facturas)

    doc_master = Document()

    for i, row_data in df_colab.iterrows():
        
        # --- TABLA 1: IDENTIFICACIÓN ENTIDAD ---
        title = doc_master.add_paragraph()
        run = title.add_run("1. IDENTIFICACIÓN DE LA ENTIDAD COLABORADORA:")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        title.paragraph_format.left_indent = Cm(-1)
        title.paragraph_format.right_indent = Cm(-0.75)
        title.paragraph_format.line_spacing = 1.0

        table = doc_master.add_table(rows=6, cols=5)
        table.style = "Table Grid"

        # Anchos columnas
        column_widths = [Cm(4.07), Cm(3.65), Cm(7.1), Cm(1.14), Cm(2.37)]
        
        # XML Tabla 1
        tbl = table._element
        tblPr = tbl.find(ns.qn('w:tblPr'))
        if tblPr is None: tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
        jc = OxmlElement('w:jc'); jc.set(ns.qn('w:val'), 'left'); tblPr.append(jc)
        tblInd = OxmlElement('w:tblInd'); tblInd.set(ns.qn('w:w'), str(int(-0.83 * 567))); tblInd.set(ns.qn('w:type'), 'dxa'); tblPr.append(tblInd)
        tblW = OxmlElement('w:tblW'); tblW.set(ns.qn('w:w'), str(int(18.33 * 567))); tblW.set(ns.qn('w:type'), 'dxa'); tblPr.append(tblW)

        for row in table.rows:
            for j, width in enumerate(column_widths): row.cells[j].width = width

        # Fusiones
        table.cell(3, 0).merge(table.cell(5, 0))
        table.cell(0, 1).merge(table.cell(0, 2))
        table.cell(1, 1).merge(table.cell(1, 2))
        table.cell(2, 1).merge(table.cell(2, 2))
        table.cell(3, 2).merge(table.cell(3, 4))
        table.cell(4, 2).merge(table.cell(4, 4))
        table.cell(5, 2).merge(table.cell(5, 4))

        fields = {
            "Razón social": (0, 0, 0, 1),
            "País de la entidad": (1, 0, 1, 1),
            "Entidad contratante": (2, 0, 2, 1),
            "NIF": (0, 3, 0, 4),
            "NIF 2": (2, 3, 2, 4), 
            "Localidad": (3, 1, 3, 2),
            "Provincia": (4, 1, 4, 2),
            "País de realización": (5, 1, 5, 2),
        }

        set_cell_format(table.cell(1,3), " ", font_size=11, bold=True)
        set_cell_format(table.cell(3, 0), "Ubicación de la Actividad principal del proyecto", font_size=10)

        for key, (r_t, c_t, _, _) in fields.items():
            if key == "NIF 2":
                set_cell_format(table.cell(r_t, c_t), "NIF ", font_size=11, bg_color="F2F2F2")
            elif key in ["Localidad", "Provincia", "País de realización"]:
                set_cell_format(table.cell(r_t, c_t), key, font_size=10)
            elif key in ["Razón social", "NIF"]:
                set_cell_format(table.cell(r_t, c_t), key, font_size=12)
            else:
                set_cell_format(table.cell(r_t, c_t), key, font_size=11)

        for key, (_, _, r_v, c_v) in fields.items():
            if key in df_colab.columns:
                val = str(row_data[key])
                cell = table.cell(r_v, c_v)
                cell.text = val
                if cell.paragraphs:
                    run = cell.paragraphs[0].runs[0]
                    run.font.name = "Arial"
                    run.font.size = Pt(10)

        doc_master.add_paragraph()

        # --- TABLA 2: JUSTIFICACIÓN ---
        t_box = doc_master.add_paragraph()
        run_box = t_box.add_run("2. JUSTIFICACIÓN Y DESCRIPCIÓN DE LA COLABORACIÓN. DETALLE DE ACTIVIDADES A REALIZAR POR LA ENTIDAD COLABORADORA EN EL PROYECTO:")
        run_box.bold = True
        run_box.font.name = "Arial"
        run_box.font.size = Pt(10)
        t_box.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        t_box.paragraph_format.left_indent = Cm(-1)
        t_box.paragraph_format.right_indent = Cm(-0.75)

        table_box = doc_master.add_table(rows=1, cols=1)
        table_box.style = 'Table Grid'
        
        tbl = table_box._element
        tblPr = tbl.find(ns.qn('w:tblPr'))
        if tblPr is None: tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
        jc = OxmlElement('w:jc'); jc.set(ns.qn('w:val'), 'left'); tblPr.append(jc)
        tblInd = OxmlElement('w:tblInd'); tblInd.set(ns.qn('w:w'), str(int(-0.83 * 567))); tblInd.set(ns.qn('w:type'), 'dxa'); tblPr.append(tblInd)
        tblW = OxmlElement('w:tblW'); tblW.set(ns.qn('w:w'), str(int(18.33 * 567))); tblW.set(ns.qn('w:type'), 'dxa'); tblPr.append(tblW)
        
        table_box.rows[0].height = Cm(2.28)
        
        cell = table_box.cell(0,0)
        cell.text = " "
        run = cell.paragraphs[0].runs[0]
        run.font.name = "Arial"
        run.font.size = Pt(10)

        doc_master.add_paragraph()

        # --- TABLA 3: COSTE ---
        title2 = doc_master.add_paragraph()
        run = title2.add_run("3.COSTE DE LA COLABORACIÓN")
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.bold = True
        title2.paragraph_format.left_indent = Cm(-1)
        title2.paragraph_format.right_indent = Cm(-0.75)

        entidad_actual = row_data.get("Razón social", "")
        facturas_entidad = df_facturas[df_facturas["Entidad"] == entidad_actual]
        
        crear_tabla_coste_colaboracion(doc_master, facturas_entidad)

        if i < len(df_colab) - 1:
            doc_master.add_page_break()

    # Fusión Ficha 2.2
    fusionar_y_guardar(doc_master, ruta_plantilla_base, ruta_salida_final)


def fusionar_y_guardar(doc_generado, ruta_plantilla, ruta_salida):
    """Función auxiliar para fusionar con plantilla y guardar."""
    if os.path.exists(ruta_plantilla):
        doc_base = Document(ruta_plantilla)
        # Sin salto de página forzado (control manual desde Word)
    else:
        print("⚠️ No se encontró plantilla base, usando generado.")
        doc_base = doc_generado
        doc_generado = None

    if doc_generado:
        # Añadir contenido generado al final del base
        for element in doc_generado.element.body:
            doc_base.element.body.append(element)
    
    # Asegurar bordes en todas las tablas del documento final
    for table in doc_base.tables: 
        table.style = "Table Grid"
    
    doc_base.save(ruta_salida)