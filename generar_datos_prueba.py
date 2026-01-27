import pandas as pd
import os
from docx import Document

# Configuración de rutas
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_DIR = os.path.join(BASE_DIR, 'inputs')

# Crear carpeta inputs si no existe
if not os.path.exists(INPUT_DIR):
    os.makedirs(INPUT_DIR)
    print(f"📁 Carpeta creada: {INPUT_DIR}")

print("--- GENERANDO DATOS DE PRUEBA COMPLETOS ---")

# ==========================================
# 1. Generar Excel Personal (Ficha 2.1)
# ==========================================
data_personal = {
    "Nombre": ["Juan", "Maria", "Carlos"],
    "Apellidos": ["Pérez Gómez", "López Díaz", "Ruiz Martín"],
    "Coste horario (€/hora)": [35.50, 45.00, 28.75],
    "Coste total (€)": [35500.00, 45000.00, 28750.00],
    "Horas totales": [1000, 1000, 1000],
    "Coste I+D (€)": [20000.00, 30000.00, 15000.00],
    "Horas I+D": [500, 700, 500],
    "Departamento": ["I+D", "Ingeniería", "Sistemas"],
    "Puesto actual": ["Jefe de Proyecto", "Analista Senior", "Desarrollador Full Stack"],
    "Titulación 1": ["Ingeniería Informática", "Ingeniería Industrial", "Grado en Matemáticas"],
    "Titulación 2": ["Máster en IA", "", "Máster en Big Data"],
    
    # --- COLUMNAS NUEVAS PARA HISTORIAL PROFESIONAL ---
    "EMPRESA 1": ["Empresa Antigua S.L.", "Consultora Big Four", "Startup Tech"],
    "PERIODO 1": ["2015-2020", "2010-2018", "2019-2022"],
    "PUESTO 1": ["Analista Junior", "Consultor Senior", "Desarrollador Junior"],
    
    "EMPRESA 2": ["", "Empresa Mediana S.A.", ""],
    "PERIODO 2": ["", "2018-2022", ""],
    "PUESTO 2": ["", "Gerente de Proyecto", ""],
    
    "EMPRESA 3": ["", "", ""],
    "PERIODO 3": ["", "", ""],
    "PUESTO 3": ["", "", ""],
    
    # --- COLUMNAS NUEVAS PARA ACTIVIDADES ---
    "Actividad 1": [
        "Fase I. Diseño de la arquitectura del software y selección de tecnologías.",
        "Análisis de requisitos funcionales y no funcionales del sistema.",
        "Desarrollo de módulos de backend en Python."
    ],
    "Actividad 2": [
        "Fase II. Implementación de algoritmos de inteligencia artificial.",
        "Coordinación del equipo técnico y validación de entregables.",
        "Integración de APIs de terceros."
    ],
    "Actividad 3": [
        "Pruebas unitarias y de integración del sistema.",
        "Diseño de prototipos de interfaz de usuario.",
        "Optimización de consultas a base de datos."
    ],
    "Actividad 4": [
        "Redacción de documentación técnica y manuales de usuario.",
        "",
        ""
    ]
}

df_personal = pd.DataFrame(data_personal)
ruta_personal = os.path.join(INPUT_DIR, "Excel_Personal_2.1.xlsx")
df_personal.to_excel(ruta_personal, index=False)
print(f"✅ Generado Excel Personal: {ruta_personal}")

# ==========================================
# 2. Generar Excel Colaboraciones (Ficha 2.2)
# ==========================================
data_colaboraciones = {
    "Razón social": ["Tech Solutions S.L.", "Consultoría Global S.A."],
    "NIF": ["B12345678", "A87654321"],
    "NIF 2": ["B12345678", "A87654321"], # Columna extra a veces requerida
    "País de la entidad": ["España", "Francia"],
    "Entidad contratante": ["Mi Empresa S.L.", "Mi Empresa S.L."],
    "Localidad": ["Madrid", "París"],
    "Provincia": ["Madrid", "Île-de-France"],
    "País de realización": ["España", "Francia"]
}

df_colab = pd.DataFrame(data_colaboraciones)
ruta_colab = os.path.join(INPUT_DIR, "Excel_Colaboraciones_2.2.xlsx")
df_colab.to_excel(ruta_colab, index=False)
print(f"✅ Generado Excel Colaboraciones: {ruta_colab}")

# ==========================================
# 3. Generar Excel Facturas (Ficha 2.2)
# ==========================================
data_facturas = {
    "Entidad": ["Tech Solutions S.L.", "Tech Solutions S.L.", "Consultoría Global S.A."],
    "Nombre factura": ["Factura F-001 Licencias Software", "Factura F-002 Soporte Técnico", "Factura C-99 Auditoría Externa"],
    "Importe (€)": [1500.50, 500.00, 3200.00]
}

df_facturas = pd.DataFrame(data_facturas)
ruta_facturas = os.path.join(INPUT_DIR, "Excel_Facturas_2.2.xlsx")
df_facturas.to_excel(ruta_facturas, index=False)
print(f"✅ Generado Excel Facturas: {ruta_facturas}")

# ==========================================
# 4. Generar Plantillas Word vacías (Dummy)
# ==========================================
def crear_plantilla_dummy(nombre):
    ruta = os.path.join(INPUT_DIR, nombre)
    if not os.path.exists(ruta):
        doc = Document()
        doc.add_paragraph("--- PORTADA PLANTILLA (DUMMY) ---")
        doc.add_paragraph("Esta página representa la carátula o instrucciones iniciales.")
        # IMPORTANTE: Un salto de página al final para que lo nuevo empiece limpio
        doc.add_page_break() 
        doc.save(ruta)
        print(f"⚠️  Plantilla dummy creada: {ruta}")
    else:
        print(f"ℹ️  La plantilla {nombre} ya existe, no se sobrescribe.")

crear_plantilla_dummy("2.1.docx")
crear_plantilla_dummy("2.2.docx")

print("\n🎉 ¡Datos de prueba generados! Ahora ejecuta 'python src/main.py'")